#!/usr/bin/env python3
"""
PDF Tool — Watermark Removal + OCR to Word
==========================================
Processes all PDF files in the current working directory.

Modes (interactive prompt at launch):
  1  — Remove watermark only          (PDF → clean PDF)
  2  — Remove watermark + OCR to Word (PDF → clean PDF → .docx)
  3  — OCR to Word only               (PDF → .docx, source unchanged)

Watermark removal strategies (in order):
  1. QQAP marker (A-PDF Watermark format)
  2. Gray-text q…Q content blocks
  3. Opacity-based Form XObjects
  4. Rasterise + pixel-level erase (image fallback)

OCR engine: Ollama glm-ocr (auto-pulled if not present)

Install:
    pip install pymupdf numpy opencv-python tqdm python-docx ollama
"""

import base64
import os
import re
import shutil
import subprocess
import sys
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import fitz          # PyMuPDF
import numpy as np
import ollama
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from tqdm import tqdm


# ═══════════════════════════════════════════════════════════════════════════════
#  WATERMARK REMOVAL  (kept exactly as provided — v3)
# ═══════════════════════════════════════════════════════════════════════════════

BLOCK_RE    = re.compile(rb'q\b(.*?)\bQ\b', re.DOTALL)
GRAY_RG_RE  = re.compile(rb'([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+(?:rg|RG)')
CA_RE       = re.compile(rb'([\d.]+)\s+ca\b')
BIG_FONT_RE = re.compile(rb'/\S+\s+([\d.]+)\s+Tf')


def _is_gray_watermark_block(block_bytes: bytes) -> bool:
    if b'BT' not in block_bytes or b'ET' not in block_bytes:
        return False
    ca_matches = CA_RE.findall(block_bytes)
    if ca_matches:
        opacity = float(ca_matches[0])
        if opacity < 0.7:
            return True
    color_matches = GRAY_RG_RE.findall(block_bytes)
    if not color_matches:
        return False
    for rm in color_matches:
        r, g, b = float(rm[0]), float(rm[1]), float(rm[2])
        if abs(r - g) < 0.08 and abs(g - b) < 0.08 and 0.35 < r < 0.90:
            return True
    return False


def _strip_watermark_blocks(stream: bytes) -> tuple[bytes, int]:
    removed = 0
    result  = bytearray()
    last_end = 0
    for m in BLOCK_RE.finditer(stream):
        if _is_gray_watermark_block(m.group(1)):
            result.extend(stream[last_end:m.start()])
            removed += 1
        else:
            result.extend(stream[last_end:m.end()])
        last_end = m.end()
    result.extend(stream[last_end:])
    return bytes(result), removed


def _stream_is_all_watermark(stream: bytes) -> bool:
    blocks = BLOCK_RE.findall(stream)
    if not blocks:
        return False
    return all(_is_gray_watermark_block(b) for b in blocks)


def _remove_watermark_from_image_array(img_bgr: np.ndarray) -> np.ndarray:
    import cv2
    gray     = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY).astype(np.float32)
    img_f    = img_bgr.astype(np.float32)
    ch_spread = img_f.max(axis=2) - img_f.min(axis=2)
    wm_mask  = ((gray > 115) & (gray < 244) & (ch_spread < 80)).astype(np.uint8)
    wm_mask[gray < 115] = 0
    result   = img_bgr.copy()
    result[wm_mask > 0] = [255, 255, 255]
    return result


def remove_watermarks_from_pdf(src: str, dst: str) -> dict:
    """Remove watermarks from src PDF and write to dst. Returns result stats."""
    doc = fitz.open(src)
    results = {
        'pages': [],
        'streams_zeroed': 0,
        'streams_stripped': 0,
        'xobjects_cleared': 0,
        'used_image_fallback': False,
    }

    for page_num, page in enumerate(doc):
        xrefs        = page.get_contents()
        page_zeroed  = 0
        page_stripped = 0
        page_xobj    = 0

        # Strategy 1 — QQAP
        qqap_indices = set()
        for idx, xref in enumerate(xrefs):
            try:
                if doc.xref_is_stream(xref) and b'QQAP' in doc.xref_stream(xref):
                    qqap_indices.add(idx)
            except Exception:
                pass
        to_zero = set(qqap_indices)
        for idx in list(qqap_indices):
            for neighbor in [idx - 1, idx + 1]:
                if 0 <= neighbor < len(xrefs):
                    try:
                        s = doc.xref_stream(xrefs[neighbor])
                        if s.strip() in (b'q', b'Q', b''):
                            to_zero.add(neighbor)
                    except Exception:
                        pass
        for idx in to_zero:
            try:
                doc.update_stream(xrefs[idx], b'')
                page_zeroed += 1
            except Exception:
                pass

        # Strategy 2 — Gray-text blocks
        xrefs = page.get_contents()
        for xref in xrefs:
            try:
                if not doc.xref_is_stream(xref):
                    continue
                stream = doc.xref_stream(xref)
                if not stream.strip():
                    continue
                if _stream_is_all_watermark(stream):
                    doc.update_stream(xref, b'')
                    page_zeroed += 1
                else:
                    cleaned, n_removed = _strip_watermark_blocks(stream)
                    if n_removed > 0 and cleaned != stream:
                        doc.update_stream(xref, cleaned)
                        page_stripped += n_removed
            except Exception:
                pass

        # Strategy 3 — Opacity Form XObjects
        for xobj_xref, name, *_ in page.get_xobjects():
            try:
                obj_dict = doc.xref_object(xobj_xref)
                is_qqap = 'QQAP' in str(name) or 'QQAP' in obj_dict
                has_transparency = ('/ca' in obj_dict or '/CA' in obj_dict
                                    or '/Transparency' in obj_dict)
                is_form = '/Form' in obj_dict
                if is_qqap or (is_form and has_transparency):
                    doc.update_stream(xobj_xref, b'')
                    page_xobj += 1
            except Exception:
                pass

        qqap_left = sum(
            1 for xref in page.get_contents()
            if doc.xref_is_stream(xref) and b'QQAP' in (doc.xref_stream(xref) or b'')
        )
        results['streams_zeroed']  += page_zeroed
        results['streams_stripped'] += page_stripped
        results['xobjects_cleared'] += page_xobj
        results['pages'].append({
            'page': page_num + 1,
            'streams_zeroed': page_zeroed,
            'streams_stripped': page_stripped,
            'xobj_cleared': page_xobj,
            'qqap_remaining': qqap_left,
            'doc_words': len(page.get_text('text').split()),
        })

    total_removed = (results['streams_zeroed'] +
                     results['streams_stripped'] +
                     results['xobjects_cleared'])

    if total_removed > 0:
        doc.save(dst, garbage=4, deflate=True)
        doc.close()
    else:
        # Strategy 4 — Image fallback
        print('    Strategies 1-3 found nothing; using image-based removal…')
        results['used_image_fallback'] = True
        doc.close()

        import cv2
        orig    = fitz.open(src)
        new_doc = fitz.open()
        tmp_dir = os.path.join(os.path.dirname(dst) or '.', '_tmp_wm')
        os.makedirs(tmp_dir, exist_ok=True)

        for page_num, page in enumerate(orig):
            mat      = fitz.Matrix(2, 2)
            pix      = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = np.frombuffer(pix.samples, dtype=np.uint8)
            img_bgr  = img_bytes.reshape(pix.height, pix.width, 3)[:, :, ::-1].copy()
            cleaned  = _remove_watermark_from_image_array(img_bgr)
            tmp_path = os.path.join(tmp_dir, f'p{page_num:04d}.png')
            cv2.imwrite(tmp_path, cleaned)
            cleaned_pix = fitz.Pixmap(tmp_path)
            new_page    = new_doc.new_page(width=page.rect.width, height=page.rect.height)
            new_page.insert_image(new_page.rect, pixmap=cleaned_pix)
            os.unlink(tmp_path)
            results['pages'][page_num]['doc_words'] = len(page.get_text('text').split())

        new_doc.save(dst, garbage=4, deflate=True)
        new_doc.close()
        orig.close()
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return results


# ═══════════════════════════════════════════════════════════════════════════════
#  OCR → WORD  (glm-ocr via Ollama)
# ═══════════════════════════════════════════════════════════════════════════════

OCR_SYSTEM_PROMPT = """You are an expert document OCR and formatting assistant.
Your task is to extract ALL text and structural content from the provided document page image and reproduce it faithfully.

Follow these rules strictly:

CONTENT EXTRACTION:
- Extract every word, number, symbol and punctuation mark exactly as it appears.
- Preserve the original reading order (top-to-bottom, left-to-right for Western text).
- Do NOT omit any content, even if it seems repetitive or boilerplate.

FORMATTING:
- Use Markdown to represent document structure:
  - # for the main document title (first occurrence only)
  - ## for section headings / major headings
  - ### for sub-headings
  - **bold** for bold text
  - *italic* for italic text
  - `code` for monospace / code snippets
  - > for block-quotes or highlighted callout boxes
  - Blank lines between paragraphs
  - Preserve indentation levels where meaningful

TABLES:
- Reproduce tables using GitHub-flavoured Markdown table syntax.
- Keep all cell content; merge cells are not possible in Markdown, so note merges in brackets e.g. [merged: 2 cols].

LISTS:
- Use - for unordered lists, 1. 2. 3. for ordered lists.
- Preserve nesting with indentation (two spaces per level).

SPECIAL ELEMENTS:
- For figures/charts write: [FIGURE: brief description]
- For images/photos write: [IMAGE: brief description]
- For page numbers / headers / footers reproduce them prefixed with <!-- and suffixed with -->.
- For mathematical formulas use LaTeX within $…$ (inline) or $$…$$ (block).
- For signatures or stamps write: [SIGNATURE] or [STAMP: text if readable]

OUTPUT:
- Output ONLY the extracted and formatted content. No preamble, no commentary.
- If the page is blank, output exactly: [BLANK PAGE]
"""


def _ensure_model(model: str) -> None:
    """Pull the Ollama model if it is not already available."""
    try:
        available = [m.model for m in ollama.list().models]
    except Exception as exc:
        print(f'\n❌  Cannot connect to Ollama: {exc}')
        print('   Make sure Ollama is running:  ollama serve')
        sys.exit(1)

    if any(m.startswith(model) for m in available):
        return

    print(f'\n⬇️   Model "{model}" not found locally — pulling now…')
    print('   (this may take a few minutes the first time)\n')
    try:
        # Stream pull progress via subprocess so the user sees it
        subprocess.run(['ollama', 'pull', model], check=True)
        print(f'   ✅  Model "{model}" ready.\n')
    except FileNotFoundError:
        print('❌  "ollama" command not found in PATH.')
        print('   Download Ollama from https://ollama.com/download')
        sys.exit(1)
    except subprocess.CalledProcessError as exc:
        print(f'❌  Failed to pull model: {exc}')
        sys.exit(1)


def _pdf_to_images(pdf_path: Path, output_dir: Path, dpi: int) -> list[Path]:
    """Render every page of a PDF to a PNG file using PyMuPDF."""
    doc  = fitz.open(str(pdf_path))
    zoom = dpi / 72.0
    mat  = fitz.Matrix(zoom, zoom)
    paths = []
    for i, page in enumerate(doc):
        img_path = output_dir / f'page_{i+1:04d}.png'
        pix = page.get_pixmap(matrix=mat, alpha=False)
        pix.save(str(img_path))
        paths.append(img_path)
    doc.close()
    return paths


def _image_to_base64(p: Path) -> str:
    return base64.b64encode(p.read_bytes()).decode()


def _ocr_page(image_path: Path, model: str) -> str:
    response = ollama.chat(
        model=model,
        messages=[
            {'role': 'system', 'content': OCR_SYSTEM_PROMPT},
            {
                'role': 'user',
                'content': 'Please extract and format all content from this document page.',
                'images': [_image_to_base64(image_path)],
            },
        ],
    )
    return response.message.content.strip()


def _ocr_page_wrapper(args):
    image_path, model, page_num = args
    try:
        return page_num, _ocr_page(image_path, model), None
    except Exception as exc:
        return page_num, '', str(exc)


def _markdown_to_docx(doc: Document, md: str, page_num: int) -> None:
    if page_num > 1:
        doc.add_page_break()

    lines     = md.split('\n')
    i         = 0
    in_table  = False
    tbl_rows: list[list[str]] = []

    def flush_table():
        nonlocal in_table, tbl_rows
        if not tbl_rows:
            in_table = False
            return
        data = [r for r in tbl_rows if not all(
            re.match(r'^:?-+:?$', c.strip()) for c in r if c.strip()
        )]
        if not data:
            in_table = False
            tbl_rows = []
            return
        cols = max(len(r) for r in data)
        t    = doc.add_table(rows=len(data), cols=cols)
        t.style = 'Table Grid'
        for ri, row in enumerate(data):
            for ci in range(cols):
                cell = row[ci].strip() if ci < len(row) else ''
                cell = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
                cell = re.sub(r'\*(.+?)\*',     r'\1', cell)
                t.cell(ri, ci).text = cell
        doc.add_paragraph()
        in_table = False
        tbl_rows = []

    def add_inline(para, text: str) -> None:
        for tok in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text):
            if tok.startswith('**') and tok.endswith('**'):
                para.add_run(tok[2:-2]).bold = True
            elif tok.startswith('*') and tok.endswith('*'):
                para.add_run(tok[1:-1]).italic = True
            elif tok.startswith('`') and tok.endswith('`'):
                r = para.add_run(tok[1:-1])
                r.font.name = 'Courier New'
            else:
                para.add_run(tok)

    while i < len(lines):
        line    = lines[i]
        stripped = line.strip()

        # HTML comment (header/footer)
        if stripped.startswith('<!--') and stripped.endswith('-->'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[4:-3].strip())
            r.font.size = Pt(9)
            i += 1
            continue

        # Table row
        if re.match(r'^\|.*\|', stripped):
            tbl_rows.append([c.strip() for c in stripped.strip('|').split('|')])
            in_table = True
            i += 1
            continue
        elif in_table:
            flush_table()

        if not stripped:
            i += 1
            continue

        # Heading
        m = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if m:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', re.sub(r'\*(.+?)\*', r'\1', m.group(2)))
            doc.add_heading(text, level=min(len(m.group(1)), 4))
            i += 1
            continue

        # Block quote
        if stripped.startswith('> '):
            doc.add_paragraph(style='Intense Quote').add_run(stripped[2:])
            i += 1
            continue

        # Ordered list
        m = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if m:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', re.sub(r'\*(.+?)\*', r'\1', m.group(3)))
            p    = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 * (len(m.group(1)) // 2))
            p.add_run(text)
            i += 1
            continue

        # Unordered list
        m = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if m:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', re.sub(r'\*(.+?)\*', r'\1', m.group(2)))
            p    = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (len(m.group(1)) // 2))
            p.add_run(text)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}$', stripped):
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Normal paragraph
        add_inline(doc.add_paragraph(), stripped)
        i += 1

    if in_table:
        flush_table()


def ocr_pdf_to_docx(pdf_path: Path, docx_path: Path,
                    model: str, dpi: int, workers: int) -> None:
    """OCR a PDF (page by page) and write a .docx file."""
    with tempfile.TemporaryDirectory(prefix='pdf_ocr_') as tmp:
        tmp_dir = Path(tmp)

        print(f'   Rendering pages at {dpi} DPI …')
        images  = _pdf_to_images(pdf_path, tmp_dir, dpi)
        n_pages = len(images)
        print(f'   {n_pages} page(s) rendered')

        ocr_args   = [(img, model, idx + 1) for idx, img in enumerate(images)]
        page_texts: dict[int, str] = {}
        errors: dict[int, str]     = {}
        eff_w = min(workers, n_pages)

        with ThreadPoolExecutor(max_workers=eff_w) as pool:
            futures = {pool.submit(_ocr_page_wrapper, a): a for a in ocr_args}
            with tqdm(total=n_pages,
                      desc=f'   OCR ({eff_w} thread{"s" if eff_w > 1 else ""})',
                      unit='page',
                      bar_format='{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}]'
                      ) as pbar:
                for future in as_completed(futures):
                    pn, text, err = future.result()
                    if err:
                        errors[pn]     = err
                        page_texts[pn] = f'[OCR ERROR on page {pn}: {err}]'
                    else:
                        page_texts[pn] = text
                    pbar.update(1)

        if errors:
            print(f'   ⚠️  OCR errors on pages: {sorted(errors.keys())}')

        print(f'   Assembling {docx_path.name} …')
        doc   = Document()
        title = doc.add_heading(
            pdf_path.stem.replace('_', ' ').replace('-', ' ').title(), 0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for pn in sorted(page_texts):
            _markdown_to_docx(doc, page_texts[pn], pn)

        doc.save(str(docx_path))


# ═══════════════════════════════════════════════════════════════════════════════
#  INTERACTIVE MENU + ORCHESTRATION
# ═══════════════════════════════════════════════════════════════════════════════

BANNER = """
╔══════════════════════════════════════════════════════════╗
║              PDF Tool  —  Watermark + OCR                ║
╠══════════════════════════════════════════════════════════╣
║  1  Remove watermark only          (PDF → clean PDF)     ║
║  2  Remove watermark + OCR to Word (PDF → clean PDF      ║
║                                         → .docx)         ║
║  3  OCR to Word only               (PDF → .docx)         ║
╚══════════════════════════════════════════════════════════╝
"""

OCR_MODES = {2, 3}   # modes that need Ollama
WM_MODES  = {1, 2}   # modes that need watermark removal


def _ask_mode() -> int:
    print(BANNER)
    while True:
        raw = input('Enter mode [1 / 2 / 3]: ').strip()
        if raw in ('1', '2', '3'):
            return int(raw)
        print('  Please enter 1, 2, or 3.')


def _ask_int(prompt: str, default: int, lo: int, hi: int) -> int:
    raw = input(f'{prompt} [{default}]: ').strip()
    if not raw:
        return default
    try:
        v = int(raw)
        if lo <= v <= hi:
            return v
    except ValueError:
        pass
    print(f'  Invalid input — using default ({default})')
    return default


def _watermark_summary(res: dict) -> str:
    method = 'image-fallback' if res['used_image_fallback'] else 'stream'
    zeroed = res['streams_zeroed']
    stripped = res['streams_stripped']
    xobj   = res['xobjects_cleared']
    return f'zeroed={zeroed} stripped={stripped} xobj={xobj} method={method}'


def main() -> None:
    # ── Mode selection ────────────────────────────────────────────────────────
    mode = _ask_mode()
    mode_labels = {
        1: 'Remove watermark only',
        2: 'Remove watermark + OCR to Word',
        3: 'OCR to Word only',
    }
    print(f'\n  Selected: [{mode}] {mode_labels[mode]}\n')

    # ── OCR settings (only when needed) ──────────────────────────────────────
    ocr_model   = 'glm-ocr'
    ocr_dpi     = 200
    ocr_workers = 4

    if mode in OCR_MODES:
        print('  OCR settings (press Enter to keep defaults):')
        ocr_dpi     = _ask_int('    DPI (100-400)',    200, 100, 400)
        ocr_workers = _ask_int('    Threads (1-16)',     4,   1,  16)
        print()

    # ── Find PDFs ─────────────────────────────────────────────────────────────
    scan_dir  = Path('.').resolve()
    pdf_files = sorted(p for p in scan_dir.glob('*.pdf')
                       if '_cleaned' not in p.stem)

    if not pdf_files:
        print(f'No PDF files found in: {scan_dir}')
        sys.exit(0)

    print(f'🔍  Found {len(pdf_files)} PDF file(s) in {scan_dir}')

    # ── Output folder for clean PDFs ──────────────────────────────────────────
    if mode in WM_MODES:
        cleaned_dir = scan_dir / 'cleaned'
        cleaned_dir.mkdir(exist_ok=True)
        print(f'📁  Cleaned PDFs → {cleaned_dir}')

    # ── Ensure Ollama model is available ─────────────────────────────────────
    if mode in OCR_MODES:
        _ensure_model(ocr_model)
        print(f'🤖  Model    : {ocr_model}')
        print(f'🖼️   DPI      : {ocr_dpi}')
        print(f'⚡  Threads  : {ocr_workers}')

    print()

    # ── Process files ─────────────────────────────────────────────────────────
    successes, failures = [], []

    overall = tqdm(
        pdf_files, desc='Overall progress', unit='file',
        bar_format='{l_bar}{bar}| {n_fmt}/{total_fmt} files'
    )

    for pdf in overall:
        overall.set_postfix(file=pdf.name[:35])
        print(f'\n📄 {pdf.name}')

        try:
            # ── Step A: Watermark removal ─────────────────────────────────────
            if mode in WM_MODES:
                clean_pdf = cleaned_dir / (pdf.stem + '_cleaned.pdf')
                print(f'   Removing watermark …')
                res = remove_watermarks_from_pdf(str(pdf), str(clean_pdf))
                print(f'   ✅  Watermark done  ({_watermark_summary(res)})')
                ocr_source = clean_pdf          # OCR from the clean version
            else:
                ocr_source = pdf                # OCR directly from source

            # ── Step B: OCR ───────────────────────────────────────────────────
            if mode in OCR_MODES:
                docx_path = pdf.with_suffix('.docx')  # always next to original
                print(f'   Running OCR → {docx_path.name} …')
                ocr_pdf_to_docx(ocr_source, docx_path,
                                model=ocr_model, dpi=ocr_dpi, workers=ocr_workers)
                print(f'   ✅  Word file saved → {docx_path.name}')

            successes.append(pdf)

        except Exception as exc:
            import traceback
            print(f'\n❌  Error processing {pdf.name}: {exc}')
            traceback.print_exc()
            failures.append((pdf, str(exc)))

    # ── Summary ───────────────────────────────────────────────────────────────
    print('\n' + '═' * 60)
    print(f'✅  Succeeded : {len(successes)} file(s)')
    if failures:
        print(f'❌  Failed    : {len(failures)} file(s)')
        for p, err in failures:
            print(f'   • {p.name}: {err}')
    print('═' * 60)


if __name__ == '__main__':
    main()